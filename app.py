from flask import Flask, render_template, request, send_file, session
import cohere
from docx import Document
from docx.shared import Inches
from xhtml2pdf import pisa
from io import BytesIO

app = Flask(__name__)
app.secret_key = 'V12345'

# Initialize Cohere client
co = cohere.Client('XJOhhwIHtY5GSkBWsbDQv7WYbHCht1ttR4zLtzs1')
  

@app.route('/', methods=['GET', 'POST'])
def resume_form():
    if request.method == 'POST':
        name = request.form.get('name')
        job_title = request.form.get('job_title')
        experience = request.form.get('experience')
        education = request.form.get('education')
        skills = request.form.get('skills')
        projects = request.form.get('projects')
        template = request.form.get('template')

        prompt = f"""
Write a complete professional ATS resume for the following candidate.

Name: {name}
Target Job Title: {job_title}
Experience: {experience}
Education: {education}
Skills: {skills}
Projects: {projects}

Format the output with the following sections:
1. Professional Summary
2. Work Experience
3. Education
4. Skills
5. Projects

Use a formal and confident tone. Keep it ATS-friendly.
"""

        response = co.generate(
            model='command',
            prompt=prompt,
            max_tokens=800,
            temperature=0.6
        )

        resume_text = response.generations[0].text.strip()

        session['resume'] = resume_text
        session['name'] = name
        session['template'] = template

        return render_template('result.html', resume=resume_text, name=name)

    return render_template('form.html')


@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    try:
        resume_text = session.get('resume', '')
        name = session.get('name', 'Resume')
        template = session.get('template', 'pdf_classic.html')

        if not resume_text:
            return "No resume data found. Please generate the resume first."

        # Example dummy structured data
        resume = {
            'name': name,
            'summary': resume_text.split('Work Experience')[0].strip() if 'Work Experience' in resume_text else resume_text,
            'experience': [{'position': 'Software Developer', 'company': 'TechCorp', 'details': ['Built web apps', 'Collaborated with team']}],
            'education': 'BSc in Computer Science, XYZ University',
            'skills': ['Python', 'Flask', 'SQL'],
            'projects': [{'title': 'AI Resume Builder', 'description': 'A tool to generate resumes using AI'}]
        }

        rendered_html = render_template(template, name=name, resume=resume)
        pdf_file = create_pdf(rendered_html)

        if not pdf_file:
            return "Failed to generate PDF", 500

        return send_file(
            pdf_file,
            download_name=f"{name.replace(' ', '_')}_resume.pdf",
            as_attachment=True,
            mimetype='application/pdf'
        )

    except Exception as e:
        app.logger.error(f"Error generating PDF: {str(e)}")
        return f"Error generating PDF: {str(e)}", 500


@app.route('/download_docx', methods=['POST'])
def download_docx():
    try:
        resume_text = session.get('resume', '')
        name = session.get('name', 'Resume')

        if not resume_text:
            return "No resume data found. Please generate the resume first."

        # Create a new Word document
        doc = Document()
        
        # Add title with name
        doc.add_heading(f"{name}'s Resume", level=1)
        
        # Convert the plain text resume into paragraphs
        for section in resume_text.split('\n\n'):
            if section.strip():  # Skip empty sections
                # Check if this looks like a section heading
                if ':' in section or section.upper() == section:
                    doc.add_heading(section, level=2)
                else:
                    doc.add_paragraph(section)
        
        # Save to a BytesIO buffer
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            download_name=f"{name.replace(' ', '_')}_resume.docx",
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error generating DOCX: {str(e)}")
        return f"Error generating DOCX: {str(e)}", 500


def create_pdf(source_html):
    result = BytesIO()
    pisa_status = pisa.CreatePDF(BytesIO(source_html.encode('utf-8')), dest=result)
    if pisa_status.err:
        return None
    result.seek(0)
    return result


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
